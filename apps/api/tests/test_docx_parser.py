"""ADR 0003 slice 1 (DOCX) — the real DOCX parser running INSIDE the isolation
harness. Extraction correctness (blocks, exact offsets, heading→section,
non-heading→None, page=None) plus the DOCX-specific defenses (decompressed-size
cap, XML entity-expansion, corrupt/empty handling).

Runs through the real subprocess harness, so — like `test_parse_harness` — it is
Linux-only (RLIMIT_AS/timeout kill; ADR 0003 marks macOS host unsupported) and is
executed in the Linux api container / CI.
"""

from __future__ import annotations

import io
import sys
import zipfile

import pytest

from app.ai.documents.docx import DOCX_CONTENT_TYPE
from app.ai.parsing.harness import parse_in_subprocess
from app.shared.exceptions.documents import DocumentParseError, DocumentTooLargeError

pytestmark = pytest.mark.skipif(
    sys.platform != "linux",
    reason="parse harness (RLIMIT_AS/timeout kill) is Linux-only; ADR 0003 marks macOS unsupported",
)


def _make_docx(build) -> bytes:  # type: ignore[no-untyped-def]
    from docx import Document

    d = Document()
    build(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


async def _parse(data: bytes, **kw: object):  # type: ignore[no-untyped-def]
    return await parse_in_subprocess(parser="docx", data=data, content_type=DOCX_CONTENT_TYPE, **kw)


# --- extraction: blocks, offsets, heading sections, page=None ----------------


@pytest.mark.asyncio
async def test_docx_extracts_blocks_with_exact_offsets_and_heading_sections() -> None:
    def build(d):  # type: ignore[no-untyped-def]
        d.add_paragraph("Intro before any heading.")
        d.add_heading("Chapter One", level=1)
        d.add_paragraph("Body under chapter one.")
        d.add_heading("Sub A", level=2)
        d.add_paragraph("Body under sub a.")

    doc = await _parse(_make_docx(build))

    # page is always None for DOCX.
    assert all(b.page is None for b in doc.blocks)
    # Offset round-trip / gap-free tiling against the parent-built full_text.
    full = doc.full_text
    pos = 0
    for b in doc.blocks:
        assert full[pos : pos + len(b.text)] == b.text
        pos += len(b.text)
    assert pos == len(full)
    # Section from EXPLICIT heading styles; a paragraph before any heading → None.
    by_text = {b.text.strip(): b.section for b in doc.blocks}
    assert by_text["Intro before any heading."] is None
    assert by_text["Chapter One"] == "Chapter One"  # the heading block labels itself
    assert by_text["Body under chapter one."] == "Chapter One"
    assert by_text["Sub A"] == "Sub A"
    assert by_text["Body under sub a."] == "Sub A"


@pytest.mark.asyncio
async def test_empty_but_valid_docx_yields_zero_blocks_no_error() -> None:
    doc = await _parse(_make_docx(lambda d: None))  # blank doc → one empty paragraph
    assert doc.blocks == []
    assert doc.full_text == ""


# --- defenses: zip bomb, entity expansion, corrupt --------------------------


@pytest.mark.asyncio
async def test_zip_bomb_rejected_by_decompressed_cap() -> None:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/document.xml", b"\x00" * (5 * 1024 * 1024))  # 5 MB uncompressed
    with pytest.raises(DocumentTooLargeError):
        await _parse(buf.getvalue(), child_env={"NEO_DOCX_MAX_DECOMPRESSED": str(1024 * 1024)})


@pytest.mark.asyncio
async def test_xml_entity_expansion_docx_rejected() -> None:
    bomb = (
        b'<?xml version="1.0"?>'
        b'<!DOCTYPE lolz [<!ENTITY lol "lol"><!ENTITY lol2 "&lol;&lol;&lol;">]>'
        b"<root>&lol2;</root>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", bomb)
    with pytest.raises(DocumentParseError):
        await _parse(buf.getvalue())


@pytest.mark.asyncio
async def test_corrupt_docx_rejected() -> None:
    with pytest.raises(DocumentParseError):
        await _parse(b"PK\x03\x04 truncated not-a-real-zip")


@pytest.mark.asyncio
async def test_valid_zip_but_not_a_docx_rejected() -> None:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("random.xml", b"<r>ok</r>")  # valid zip + xml, but no OPC/document.xml
    with pytest.raises(DocumentParseError):
        await _parse(buf.getvalue())
